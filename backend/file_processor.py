import os
import uuid
import logging
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document
from config import (
    logger,
    log_with_extra,
    UPLOAD_DIR,
    ALLOWED_FILE_EXTENSIONS,
    MAX_FILE_SIZE,
)


class FileProcessor:
    def __init__(self, filename: str, file_content: bytes):
        self.filename = filename
        self.file_content = file_content
        self.file_id = str(uuid.uuid4())
        self.file_path = os.path.join(UPLOAD_DIR, f"{self.file_id}.tmp")

    @log_with_extra(stage="validation")
    def validate_file(self) -> None:
        file_ext = os.path.splitext(self.filename)[1].lower()
        if file_ext not in ALLOWED_FILE_EXTENSIONS:
            raise ValueError(
                f"Invalid file type. Allowed types: {', '.join(ALLOWED_FILE_EXTENSIONS)}"
            )

        if len(self.file_content) > MAX_FILE_SIZE:
            raise ValueError(f"File size exceeds {MAX_FILE_SIZE / (1024 * 1024)}MB limit")

        logger.info(
            "File validated",
            extra={"extra_data": {"file_id": self.file_id, "filename": self.filename, "size_bytes": len(self.file_content)}},
        )

    @log_with_extra(stage="save")
    def save_file(self) -> None:
        with open(self.file_path, "wb") as f:
            f.write(self.file_content)
        logger.info(
            "File saved",
            extra={"extra_data": {"file_id": self.file_id, "path": self.file_path}},
        )

    @log_with_extra(stage="extract_pdf")
    def extract_pdf_text(self) -> str:
        try:
            reader = PdfReader(self.file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"

            logger.info(
                "PDF text extracted",
                extra={
                    "extra_data": {
                        "file_id": self.file_id,
                        "pages": len(reader.pages),
                        "text_length": len(text),
                    }
                },
            )
            return text.strip()
        except Exception as e:
            logger.error(
                f"Failed to extract PDF text: {str(e)}",
                extra={"extra_data": {"file_id": self.file_id, "error": str(e)}},
            )
            raise

    @log_with_extra(stage="extract_docx")
    def extract_docx_text(self) -> str:
        try:
            doc = Document(self.file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            logger.info(
                "DOCX text extracted",
                extra={
                    "extra_data": {
                        "file_id": self.file_id,
                        "paragraphs": len(doc.paragraphs),
                        "text_length": len(text),
                    }
                },
            )
            return text.strip()
        except Exception as e:
            logger.error(
                f"Failed to extract DOCX text: {str(e)}",
                extra={"extra_data": {"file_id": self.file_id, "error": str(e)}},
            )
            raise

    @log_with_extra(stage="process")
    def process(self) -> str:
        self.validate_file()
        self.save_file()

        file_ext = os.path.splitext(self.filename)[1].lower()

        if file_ext == ".pdf":
            text = self.extract_pdf_text()
        elif file_ext == ".docx":
            text = self.extract_docx_text()
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        if not text or len(text.strip()) < 10:
            raise ValueError("Extracted text is too short or empty")

        logger.info(
            "File processing completed",
            extra={
                "extra_data": {
                    "file_id": self.file_id,
                    "filename": self.filename,
                    "extracted_length": len(text),
                }
            },
        )

        return text

    def cleanup(self) -> None:
        try:
            if os.path.exists(self.file_path):
                os.remove(self.file_path)
                logger.info(
                    "File cleaned up",
                    extra={"extra_data": {"file_id": self.file_id, "path": self.file_path}},
                )
        except Exception as e:
            logger.error(
                f"Failed to cleanup file: {str(e)}",
                extra={"extra_data": {"file_id": self.file_id, "error": str(e)}},
            )
